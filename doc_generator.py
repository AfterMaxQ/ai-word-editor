import json
import sys
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from numpy.distutils.conv_template import header
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ai_parser import parse_natural_language_to_json


def load_document_data(filepath):
    """
        从指定的JSON文件中读取文档结构数据。
        Args:
            filepath (str): JSON文件的路径。

        Returns:
            dict: 包含文档结构数据的字典。
                  如果文件不存在或格式错误，则程序会打印错误信息并退出。
    """
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            data = json.load(f)
        return data
    except FileNotFoundError:
        print(f"错误：文件未找到")
    except json.JSONDecodeError:
        print(f"错误：JSON文件格式不正确 -> {filepath}")
        sys.exit(1)


def apply_paragraph_properties(paragraph, properties: dict):
    """
        将properties字典中定义的格式应用到段落对象上。

        Args:
            paragraph: python-docx的段落对象。
            properties (dict): 包含格式定义的字典。
    """
    p_format = paragraph.paragraph_format

    # 设置首行缩进
    if 'first_line_indent' in properties:
        p_format.first_line_indent = Cm(properties['first_line_indent'])

    # --- 设置字体格式 (字体、字号等), 字体格式需要应用到段落内的Run上。---
    if paragraph.runs:
        font = paragraph.runs[0].font
        # 设置字体名称
        if 'font_name' in properties:
            font.name = properties['font_name']
            # 导入中文字体所需的包
            from docx.oxml.ns import qn
            # 设置中文字体 (东亚字体)
            font.element.rPr.rFonts.set(qn('w:eastAsia'), properties['font_name'])

        # 设置字体大小
        if "font_size" in properties:
            font.size = Pt(properties['font_size'])
        #设置粗体
        if 'bold' in properties:
            font.bold = bool(properties['bold'])

def add_table_from_data(doc, element: dict):
    """
        根据element字典中的数据，在文档中添加一个表格。

        Args:
            doc: The python-docx Document object.
            element (dict): 包含表格数据的字典。
    """
    properties = element.get("properties",{})
    table_data = element.get('data',[])

    # 1. 获取表格尺寸和验证
    rows = properties.get('rows', 0)
    cols = properties.get('cols', 0)
    if rows==0 or cols==0 or not table_data:
        print("警告：表格数据不完整，跳过此表格。")
        return

    # 定义一个从字符串到docx枚举的映射字典
    ALIGNMENT_MAP = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
    }
    # 获取JSON中定义的对齐方式列表
    alignments = properties.get("alignments", [])

    # 2. 创建表格
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'

    # 3. 填充数据并设置格式
    for i in range(rows):
        for j in range(cols):
            # 获取单元格对象
            cell = table.cell(i, j)
            # 填充文本 (增加边界检查，防止数据行/列数与定义的rows/cols不匹配)
            if i < len(table_data) and j < len(table_data[i]):
                cell.text = str(table_data[i][j])
            # 4. 如果是表头行，则加粗
            if properties.get('header') and i == 0:
                #单元格内的第一个段落的第一个run的字体
                if cell.paragraphs and cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].font.bold = True

            if j < len(alignments):
                align_str = alignments[j]
                alignment_enum = ALIGNMENT_MAP.get(align_str.lower())
                if alignment_enum is not None and cell.paragraphs:
                    cell.paragraphs[0].paragraph_format.alignment = alignment_enum


def create_document(data: dict):
    """
        根据传入的数据字典，创建一个Word文档对象。

        Args:
            data (dict): 从JSON文件加载的文档结构数据。

        Returns:
            Document: 一个构建好的python-docx的Document对象。
    """
    doc = Document()

    if 'elements' not in data or not isinstance(data['elements'], list):
        print("错误：JSON数据中缺少'elements'列表。")
        return doc # 返回一个空文档
    for element in data['elements']:
        element_type = element.get('type')

        if element.get('type')=='paragraph':
            text = element.get('text', '')
            properties = element.get('properties', {})
            if not isinstance(properties, dict):
                properties = {}

            style = properties.get("style")

            if style and 'Heading' in style:
                p = doc.add_paragraph(text, style=style)
            else:
                p = doc.add_paragraph(text)
                apply_paragraph_properties(p, properties)

        elif element_type == "table":
            add_table_from_data(doc, element)

    return doc

def main():
    """
    脚本的主执行函数。
    负责接收用户指令、调用AI解析、创建文档并保存。
    """
    # 1. 定义用户的自然语言指令
    user_command = """
    给我一个一级标题叫'月度销售报告'。
    然后另起一段，内容是'以下是本月的销售数据汇总：'。
    接下来，创建一个3行3列的表格，包含表头，列对齐方式是左、中、中。
    表格内容是：
    销售员, 销售额(万), 区域
    张三, 120, 华北
    李四, 98, 华东
    最后，再来一段，内容是'报告结束。'，设置为加粗。
    """

    # 2. 调用AI解析器，将自然语言转换为结构化数据
    document_data = parse_natural_language_to_json(user_command)

    # 如果解析失败，则退出
    if not document_data:
        print("文档生成失败，因为AI解析步骤出错。")
        return

    # 3. 创建文档 (这部分完全复用我们之前的成果！)
    print("\n📄 正在根据AI生成的数据结构创建Word文档...")
    document_object = create_document(document_data)
    print("✅ 成功创建Word文档对象！")

    # 4. 保存文档
    output_filename = 'final_report.docx'
    document_object.save(output_filename)
    print(f"✅ 成功将文档保存为 '{output_filename}'！")

if __name__ == "__main__":
    main()
