# src/doc_generator.py
import io
import json
import os
import re
import shutil
import sys
import tempfile
import uuid
import zipfile
import httpx
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from lxml import etree
from . import numbering_generator
from .latex_converter import latex_to_omml
import xml.etree.ElementTree as ET

# 定义模型和API URL常量
MODEL_NAME = "phi3:mini"
LATEX_PROMPT_FILE = "prompts/prompt_for_latex_convert.txt"
OLLAMA_API_URL = "http://localhost:11434/api/chat"


async def translate_latex_to_omml_llm(latex_string: str) -> str | None:
    """
    使用大语言模型将LaTeX公式字符串转换为OMML XML字符串。

    Args:
        latex_string (str): 需要转换的LaTeX公式。

    Returns:
        str | None: 转换后的OMML XML字符串，如果失败则返回None。
    """
    print(f"\n[AI-LATEX] 开始调用LLM转译LaTeX公式: '{latex_string}'")
    try:
        with open(LATEX_PROMPT_FILE, 'r', encoding='utf-8') as f:
            system_prompt = f.read()
    except FileNotFoundError:
        print(f"❌ 致命错误: 无法找到LaTeX转换提示词文件 -> {LATEX_PROMPT_FILE}")
        return None

    user_prompt = f"Convert the following LaTeX formula into a centered OMML `<m:oMathPara>` XML block.\nLaTeX Input: `{latex_string}`\nAlignment: center"

    payload = {
        "model": MODEL_NAME,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ],
        "stream": False
    }

    try:
        async with httpx.AsyncClient(timeout=60.0) as client:
            response = await client.post(OLLAMA_API_URL, json=payload)
            response.raise_for_status()

        response_data = response.json()
        omml_xml_string = response_data.get('message', {}).get('content')

        if not omml_xml_string:
            return None

        # 清理模型可能返回的 markdown 代码块
        omml_xml_string = re.sub(r'^```xml\s*|\s*```$', '', omml_xml_string, flags=re.MULTILINE).strip()

        # 验证XML的有效性
        etree.fromstring(omml_xml_string)

        print(f"  [成功] LaTeX 到 OMML 转译成功。")
        return omml_xml_string
    except (httpx.RequestError, httpx.HTTPStatusError, etree.XMLSyntaxError) as e:
        print(f"  [错误] LaTeX 到 OMML 转译失败: {e}")
        return None

ALIGNMENT_MAP = {
    'left': WD_ALIGN_PARAGRAPH.LEFT,
    'center': WD_ALIGN_PARAGRAPH.CENTER,
    'right': WD_ALIGN_PARAGRAPH.RIGHT,
}


def load_document_data(filepath):
    """
    从指定的JSON文件中读取文档结构数据。

    Args:
        filepath (str): JSON文件的路径。

    Returns:
        dict: 包含文档结构数据的字典。
    """
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            data = json.load(f)
        return data
    except FileNotFoundError:
        print(f"错误：文件未找到")
    except json.JSONDecodeError:
        print(f"错误：JSON文件格式不正确 -> {filepath}")
        sys.exit(1)


def apply_paragraph_properties(paragraph, properties: dict):
    """
    【已重构】将properties字典中定义的格式应用到段落的所有run上。
    增加了对不存在样式的容错处理，并修正了字体属性的应用逻辑。
    """
    p_format = paragraph.paragraph_format

    alignment_str = properties.get('alignment')
    if alignment_str and alignment_str in ALIGNMENT_MAP:
        p_format.alignment = ALIGNMENT_MAP[alignment_str]

    spacing_before_pt = properties.get('spacing_before')
    if spacing_before_pt is not None:
        p_format.space_before = Pt(spacing_before_pt)

    spacing_after_pt = properties.get('spacing_after')
    if spacing_after_pt is not None:
        p_format.space_after = Pt(spacing_after_pt)

    line_spacing_val = properties.get('line_spacing')
    if line_spacing_val is not None:
        p_format.line_spacing = line_spacing_val

    indent_cm = properties.get('first_line_indent')
    if indent_cm is not None:
        p_format.first_line_indent = Cm(indent_cm)

    # 确保至少有一个run来应用格式
    if not paragraph.runs:
        paragraph.add_run()

    # 将字体格式应用到所有run
    for run in paragraph.runs:
        font = run.font
        font_name = properties.get('font_name')
        if font_name:
            font.name = font_name
            # 设置东亚字体以确保中文字体生效
            rpr = run._r.get_or_add_rPr()
            rFonts = rpr.get_or_add_rFonts()
            rFonts.set(qn('w:eastAsia'), font_name)

        font_size_pt = properties.get('font_size')
        if font_size_pt is not None:
            font.size = Pt(font_size_pt)

        is_bold = properties.get('bold')
        if is_bold is not None:
            font.bold = bool(is_bold)

        font_color_hex = properties.get('font_color')
        if font_color_hex:
            if font_color_hex.startswith('#'):
                font_color_hex = font_color_hex[1:]
            try:
                font.color.rgb = RGBColor.from_string(font_color_hex)
            except ValueError:
                print(f"警告：无效的颜色格式码 '{font_color_hex}'，已跳过颜色设置。")


def _set_three_line_table_borders(table):
    """
    通过操作XML属性，为表格应用经典的三线表样式。

    Args:
        table: python-docx的Table对象。
    """
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border_el = OxmlElement(f'w:{border_name}')
        border_el.set(qn('w:val'), 'none')
        border_el.set(qn('w:sz'), '0')
        tblBorders.append(border_el)
    tblPr.append(tblBorders)

    for border_name in ['top', 'bottom']:
        border_el = tblBorders.find(qn(f'w:{border_name}'))
        if border_el is not None:
            border_el.set(qn('w:val'), 'single')
            border_el.set(qn('w:sz'), '4')

    header_row = table.rows[0]
    for cell in header_row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement("w:tcBorders")
            tcPr.append(tcBorders)

        bottom_border = OxmlElement("w:bottom")
        bottom_border.set(qn("w:val"), "single")
        bottom_border.set(qn("w:sz"), "4")
        tcBorders.append(bottom_border)


def add_table_from_data(doc, element: dict):
    """
    根据element字典中的数据，在文档中添加一个表格。

    Args:
        doc: python-docx的Document对象。
        element (dict): 包含表格数据的字典。
    """
    properties = element.get("properties", {})
    table_data = element.get('data', [])

    if not table_data or not table_data[0]:
        print("警告：表格数据为空或格式不正确，跳过此表格。")
        return

    cols = len(table_data[0])
    table_style = properties.get('style')
    is_three_line_style = table_style in ["Plain Table 1", "Plain Table 3"]

    if is_three_line_style:
        table = doc.add_table(rows=1, cols=cols)
        table.style = 'Table Grid'
    else:
        try:
            table = doc.add_table(rows=1, cols=cols, style=table_style or 'Table Grid')
        except KeyError:
            print(f"警告：找不到名为 '{table_style}' 的表格样式，已使用默认样式 'Table Grid'。")
            table = doc.add_table(rows=1, cols=cols, style='Table Grid')

    header_cells = table.rows[0].cells
    for j in range(cols):
        header_cells[j].text = str(table_data[0][j])
        if properties.get('header'):
            for run in header_cells[j].paragraphs[0].runs:
                run.font.bold = True

    for i in range(1, len(table_data)):
        row_cells = table.add_row().cells
        for j in range(cols):
            row_cells[j].text = str(table_data[i][j])

    if is_three_line_style:
        _set_three_line_table_borders(table)

    alignments = properties.get("alignments", [])
    if alignments:
        for col_idx, align_str in enumerate(alignments):
            if col_idx < cols:
                alignment_enum = ALIGNMENT_MAP.get(align_str.lower())
                if alignment_enum:
                    for row in table.rows:
                        row.cells[col_idx].paragraphs[0].paragraph_format.alignment = alignment_enum


def add_list_from_data(doc, element: dict):
    """
    根据element字典中的数据，在文档中添加一个有序或无序列表。

    Args:
        doc: python-docx的Document对象。
        element (dict): 包含列表数据的字典。
    """
    properties = element.get('properties', {})
    items = element.get('items', [])
    if not isinstance(items, list) or not items:
        print("警告：列表数据为空或格式不正确，跳过此列表。")
        return
    is_ordered = properties.get('ordered', False)
    style = "List Number" if is_ordered else 'List Bullet'
    for item_text in items:
        doc.add_paragraph(str(item_text), style=style)


def add_image_from_data(doc, element: dict):
    """
    根据element字典中的数据，在文档中添加一张图片。

    Args:
        doc: python-docx的Document对象。
        element (dict): 包含图片数据的字典。
    """
    properties = element.get('properties', {})
    path = properties.get('path')

    if not path:
        print("警告：图片元素缺少'path'属性，跳过此图片。")
        return
    width_cm = properties.get('width')
    height_cm = properties.get('height')
    width = Cm(width_cm) if width_cm is not None else None
    height = Cm(height_cm) if height_cm is not None else None

    try:
        doc.add_picture(path, width=width, height=height)
    except FileNotFoundError:
        print(f"警告：图片文件未找到 -> {path} ，跳过此图片。")
    except Exception as e:
        print(f"警告：插入图片时发生错误 -> {path} ({e})，跳过此图片。")


def add_page_number(paragraph):
    """
    在一个段落中添加页码域。

    Args:
        paragraph: python-docx的Paragraph对象。
    """
    run = paragraph.add_run()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_end)


def add_header_from_data(doc, element: dict, section):
    """
    根据element字典中的数据，在文档节中添加页眉。

    Args:
        doc: python-docx的Document对象。
        element (dict): 包含页眉数据的字典。
        section: python-docx的Section对象。
    """
    properties = element.get("properties", {})
    text = properties.get("text", "")
    align_str = properties.get('alignment', 'center')

    header = section.header
    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    paragraph.clear()
    paragraph.alignment = ALIGNMENT_MAP.get(align_str.lower(), WD_ALIGN_PARAGRAPH.CENTER)

    if '{PAGE_NUM}' in text:
        parts = text.split('{PAGE_NUM}')
        if parts[0]:
            paragraph.add_run(parts[0])
        add_page_number(paragraph)
        if len(parts) > 1 and parts[1]:
            paragraph.add_run(parts[1])
    else:
        paragraph.add_run(text)


def add_footer_from_data(doc, element: dict, section):
    """
    根据element字典中的数据，在文档节中添加页脚。

    Args:
        doc: python-docx的Document对象。
        element (dict): 包含页脚数据的字典。
        section: python-docx的Section对象。
    """
    properties = element.get("properties", {})
    text = properties.get("text", "")
    align_str = properties.get('alignment', 'center')

    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.clear()
    paragraph.alignment = ALIGNMENT_MAP.get(align_str.lower(), WD_ALIGN_PARAGRAPH.CENTER)

    if '{PAGE_NUM}' in text:
        parts = text.split('{PAGE_NUM}')
        if parts[0]:
            paragraph.add_run(parts[0])
        add_page_number(paragraph)
        if len(parts) > 1 and parts[1]:
            paragraph.add_run(parts[1])
    else:
        paragraph.add_run(text)


def post_process_footnotes(docx_bytes, footnotes_to_inject, reference_format: str = "[#]"):
    """
    【已重构 v3】通过解压、进行树感知的XML操作、再重新打包的方式，为DOCX文件添加脚注。
    此版本使用 lxml 的原生方法来创建和处理带命名空间的 XML 元素，修复了因 lxml 和
    python-docx API 不匹配导致的 TypeError，并确保生成的 OOXML 结构始终有效。

    Args:
        docx_bytes (bytes): 包含占位符的原始DOCX文件字节流。
        footnotes_to_inject (dict): 键是占位符，值是脚注文本的字典。
        reference_format (str): 引用标记的格式，'#'是数字占位符。

    Returns:
        bytes: 处理完脚注后的DOCX文件字节流。
    """
    if not footnotes_to_inject:
        return docx_bytes

    print("\n--- 开始脚注XML后处理阶段 (v3 - API Corrected) ---")
    print(f"  > 使用引用格式: '{reference_format}'")

    ns = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'rel': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        'content_types': 'http://schemas.openxmlformats.org/package/2006/content-types'
    }

    # 辅助函数，用于为 lxml 创建正确的命名空间限定标签
    def w_qn(tag):
        return f"{{{ns['w']}}}{tag}"

    temp_dir = tempfile.mkdtemp()
    try:
        with zipfile.ZipFile(io.BytesIO(docx_bytes), 'r') as zip_ref:
            zip_ref.extractall(temp_dir)

        # 定义文件路径
        document_path = os.path.join(temp_dir, 'word', 'document.xml')
        footnotes_path = os.path.join(temp_dir, 'word', 'footnotes.xml')
        rels_path = os.path.join(temp_dir, 'word', '_rels', 'document.xml.rels')
        content_types_path = os.path.join(temp_dir, '[Content_Types].xml')

        parser = etree.XMLParser(remove_blank_text=True)

        # 1. 准备或创建 footnotes.xml
        if os.path.exists(footnotes_path):
            footnotes_tree = etree.parse(footnotes_path, parser)
            footnotes_root = footnotes_tree.getroot()
        else:
            # 使用 lxml 的原生命名空间处理方式创建根元素
            footnotes_root = etree.Element(w_qn('footnotes'), nsmap=ns)
            etree.SubElement(footnotes_root, w_qn('footnote'), {w_qn('id'): '-1', w_qn('type'): 'separator'})
            etree.SubElement(footnotes_root, w_qn('footnote'), {w_qn('id'): '0', w_qn('type'): 'continuationSeparator'})
            footnotes_tree = etree.ElementTree(footnotes_root)

        # 2. 解析 document.xml 以便进行操作
        doc_tree = etree.parse(document_path, parser)
        doc_root = doc_tree.getroot()

        for placeholder, footnote_text in footnotes_to_inject.items():
            # 3. 在 footnotes.xml 中创建新的脚注内容
            existing_ids = [int(fn.get(w_qn('id'))) for fn in
                            footnotes_root.xpath('.//w:footnote[@w:id > "0"]', namespaces=ns)]
            new_footnote_id = (max(existing_ids) + 1) if existing_ids else 1

            footnote_element = etree.SubElement(footnotes_root, w_qn('footnote'), {w_qn('id'): str(new_footnote_id)})
            p_el = etree.SubElement(footnote_element, w_qn('p'))
            pPr = etree.SubElement(p_el, w_qn('pPr'))
            etree.SubElement(pPr, w_qn('pStyle'), {w_qn('val'): 'FootnoteText'})
            r_ref = etree.SubElement(p_el, w_qn('r'))
            rPr_ref = etree.SubElement(r_ref, w_qn('rPr'))
            etree.SubElement(rPr_ref, w_qn('rStyle'), {w_qn('val'): 'FootnoteReference'})
            etree.SubElement(r_ref, w_qn('footnoteRef'))
            r_text = etree.SubElement(p_el, w_qn('r'))
            t_text = etree.SubElement(r_text, w_qn('t'))
            t_text.text = f" {footnote_text}"
            # qn() for 'xml:space' is fine because it's a different, standard namespace prefix
            t_text.set(qn('xml:space'), 'preserve')

            # 4. 在 document.xml 中查找并替换占位符
            target_t_elements = doc_root.xpath(f".//w:t[text()='{placeholder}']", namespaces=ns)

            if not target_t_elements:
                print(f"警告: 在 document.xml 中未找到占位符 '{placeholder}'。")
                continue

            target_t = target_t_elements[0]
            run_to_replace = target_t.getparent()
            parent_paragraph = run_to_replace.getparent()

            # 5. 创建新的脚注引用 run 序列
            parts = reference_format.split('#')
            prefix, suffix = (parts[0], parts[1]) if len(parts) > 1 else (parts[0], "")

            def create_styled_run(text_content=None):
                run = etree.Element(w_qn('r'))
                rPr = etree.SubElement(run, w_qn('rPr'))
                etree.SubElement(rPr, w_qn('vertAlign'), {w_qn('val'): 'superscript'})
                if text_content:
                    t = etree.SubElement(run, w_qn('t'))
                    t.text = text_content
                return run

            new_runs = []
            if prefix: new_runs.append(create_styled_run(prefix))

            ref_run = create_styled_run()
            etree.SubElement(ref_run, w_qn('footnoteReference'), {w_qn('id'): str(new_footnote_id)})
            new_runs.append(ref_run)

            if suffix: new_runs.append(create_styled_run(suffix))

            # 6. 执行树操作：在占位符run之前插入新run，然后删除占位符run
            # 倒序插入以保持正确顺序
            for r in reversed(new_runs):
                run_to_replace.addprevious(r)
            parent_paragraph.remove(run_to_replace)
            print(f"  > 找到并替换脚注占位符: '{placeholder}' (Tree Method)")

        # 7. 更新关系和内容类型文件 (如果需要)
        rels_tree = etree.parse(rels_path, parser)
        rels_root = rels_tree.getroot()
        if not any(rel.get('Target') == 'footnotes.xml' for rel in rels_root):
            existing_rids = [int(r.get('Id')[3:]) for r in rels_root]
            new_rid = f"rId{max(existing_rids) + 1 if existing_rids else 1}"
            etree.SubElement(rels_root, etree.QName(ns['rel'], 'Relationship'), {
                'Id': new_rid,
                'Type': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes',
                'Target': 'footnotes.xml'
            })
            rels_tree.write(rels_path, encoding='UTF-8', xml_declaration=True, pretty_print=True)

        types_tree = etree.parse(content_types_path, parser)
        types_root = types_tree.getroot()
        if not any(ov.get('PartName') == '/word/footnotes.xml' for ov in types_root):
            etree.SubElement(types_root, etree.QName(ns['content_types'], 'Override'), {
                'PartName': '/word/footnotes.xml',
                'ContentType': 'application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml'
            })
            types_tree.write(content_types_path, encoding='UTF-8', xml_declaration=True, pretty_print=True)

        # 8. 将所有修改后的文件写回磁盘
        footnotes_tree.write(footnotes_path, encoding='UTF-8', xml_declaration=True, pretty_print=True)
        doc_tree.write(document_path, encoding='UTF-8', xml_declaration=True, pretty_print=True)

        # 9. 将整个目录重新打包成 .docx 文件的字节流
        output_stream = io.BytesIO()
        with zipfile.ZipFile(output_stream, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for root, _, files in os.walk(temp_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    archive_name = os.path.relpath(file_path, temp_dir)
                    zipf.write(file_path, archive_name)

        print("--- 脚注XML后处理阶段完成 ---")
        return output_stream.getvalue()

    finally:
        shutil.rmtree(temp_dir)


def post_process_endnotes(docx_bytes, endnotes_to_inject, reference_format: str = "#"):
    """
    通过解压、操作XML、再重新打包的方式，为DOCX文件添加尾注。

    Args:
        docx_bytes (bytes): 包含占位符的原始DOCX文件字节流。
        endnotes_to_inject (dict): 键是占位符，值是尾注文本的字典。
        reference_format (str): 引用标记的格式，'#'是数字占位符。

    Returns:
        bytes: 处理完尾注后的DOCX文件字节流。
    """
    if not endnotes_to_inject:
        return docx_bytes

    print("\n--- 开始尾注XML后处理阶段 ---")
    print(f"  > 使用引用格式: '{reference_format}'")

    ns = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'rel': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        'content_types': 'http://schemas.openxmlformats.org/package/2006/content-types'
    }
    for prefix, uri in ns.items():
        ET.register_namespace(prefix, uri)

    temp_dir = tempfile.mkdtemp()
    try:
        with zipfile.ZipFile(io.BytesIO(docx_bytes), 'r') as zip_ref:
            zip_ref.extractall(temp_dir)

        document_path = os.path.join(temp_dir, 'word', 'document.xml')
        endnotes_path = os.path.join(temp_dir, 'word', 'endnotes.xml')
        rels_path = os.path.join(temp_dir, 'word', '_rels', 'document.xml.rels')
        content_types_path = os.path.join(temp_dir, '[Content_Types].xml')

        if os.path.exists(endnotes_path):
            endnotes_tree = ET.parse(endnotes_path)
            endnotes_root = endnotes_tree.getroot()
        else:
            endnotes_root = ET.Element(f'{{{ns["w"]}}}endnotes')
            ET.SubElement(endnotes_root, f'{{{ns["w"]}}}endnote',
                          {f'{{{ns["w"]}}}id': '-1', f'{{{ns["w"]}}}type': 'separator'})
            ET.SubElement(endnotes_root, f'{{{ns["w"]}}}endnote',
                          {f'{{{ns["w"]}}}id': '0', f'{{{ns["w"]}}}type': 'continuationSeparator'})
            endnotes_tree = ET.ElementTree(endnotes_root)

        with open(document_path, 'r', encoding='utf-8') as f:
            doc_xml_content = f.read()

        for placeholder, endnote_text in endnotes_to_inject.items():
            existing_ids = [int(en.get(f'{{{ns["w"]}}}id')) for en in endnotes_root.findall('w:endnote', ns) if
                            int(en.get(f'{{{ns["w"]}}}id')) > 0]
            new_endnote_id = (max(existing_ids) + 1) if existing_ids else 1

            endnote_element = ET.SubElement(endnotes_root, f'{{{ns["w"]}}}endnote',
                                            {f'{{{ns["w"]}}}id': str(new_endnote_id)})
            p_element = ET.SubElement(endnote_element, f'{{{ns["w"]}}}p')
            pPr = ET.SubElement(p_element, f'{{{ns["w"]}}}pPr')
            ET.SubElement(pPr, f'{{{ns["w"]}}}pStyle', {f'{{{ns["w"]}}}val': 'EndnoteText'})
            r_ref = ET.SubElement(p_element, f'{{{ns["w"]}}}r')
            rPr_ref = ET.SubElement(r_ref, f'{{{ns["w"]}}}rPr')
            ET.SubElement(rPr_ref, f'{{{ns["w"]}}}rStyle', {f'{{{ns["w"]}}}val': 'EndnoteReference'})
            ET.SubElement(r_ref, f'{{{ns["w"]}}}endnoteRef')
            r_text = ET.SubElement(p_element, f'{{{ns["w"]}}}r')
            t_text = ET.SubElement(r_text, f'{{{ns["w"]}}}t')
            t_text.text = f" {endnote_text}"

            parts = reference_format.split('#')
            prefix = parts[0]
            suffix = parts[1] if len(parts) > 1 else ""

            def create_styled_run(text_content=None):
                run = ET.Element(f'{{{ns["w"]}}}r')
                rpr = ET.SubElement(run, f'{{{ns["w"]}}}rPr')
                ET.SubElement(rpr, f'{{{ns["w"]}}}vertAlign', {f'{{{ns["w"]}}}val': 'superscript'})
                if text_content:
                    t = ET.SubElement(run, f'{{{ns["w"]}}}t')
                    t.text = text_content
                return run

            xml_parts_to_combine = []
            if prefix:
                xml_parts_to_combine.append(create_styled_run(prefix))

            ref_run = create_styled_run()
            ET.SubElement(ref_run, f'{{{ns["w"]}}}endnoteReference', {f'{{{ns["w"]}}}id': str(new_endnote_id)})
            xml_parts_to_combine.append(ref_run)

            if suffix:
                xml_parts_to_combine.append(create_styled_run(suffix))

            combined_xml_str = "".join([ET.tostring(part, encoding='unicode') for part in xml_parts_to_combine])

            if placeholder in doc_xml_content:
                doc_xml_content = doc_xml_content.replace(placeholder, combined_xml_str)
                print(f"  > 找到并替换尾注占位符: '{placeholder}'")
            else:
                print(f"警告: 在 document.xml 中未找到尾注占位符 '{placeholder}'。")

        endnotes_tree.write(endnotes_path, encoding='UTF-8', xml_declaration=True)
        with open(document_path, 'w', encoding='utf-8') as f:
            f.write(doc_xml_content)

        rels_tree = ET.parse(rels_path)
        rels_root = rels_tree.getroot()
        if not any(rel.get('Target') == 'endnotes.xml' for rel in rels_root):
            existing_rids = [int(r.get('Id')[3:]) for r in rels_root]
            new_rid = f"rId{max(existing_rids) + 1 if existing_rids else 1}"
            ET.SubElement(rels_root, f'{{{ns["rel"]}}}Relationship', {
                'Id': new_rid,
                'Type': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes',
                'Target': 'endnotes.xml'
            })
            rels_tree.write(rels_path, encoding='UTF-8', xml_declaration=True)

        types_tree = ET.parse(content_types_path)
        types_root = types_tree.getroot()
        if not any(
                ov.get('PartName') == '/word/endnotes.xml' for ov in types_root.findall('content_types:Override', ns)):
            ET.SubElement(types_root, f'{{{ns["content_types"]}}}Override', {
                'PartName': '/word/endnotes.xml',
                'ContentType': 'application/vnd.openxmlformats-officedocument.wordprocessingml.endnotes+xml'
            })
            types_tree.write(content_types_path, encoding='UTF-8', xml_declaration=True)

        output_stream = io.BytesIO()
        with zipfile.ZipFile(output_stream, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for root, _, files in os.walk(temp_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    archive_name = os.path.relpath(file_path, temp_dir)
                    zipf.write(file_path, archive_name)

        print("--- 尾注XML后处理阶段完成 ---")
        return output_stream.getvalue()

    finally:
        shutil.rmtree(temp_dir)


_bookmark_id_counter = 0


def add_bookmark(paragraph, bookmark_name: str):
    """
    为一个段落的完整内容添加书签。

    Args:
        paragraph: python-docx的Paragraph对象。
        bookmark_name (str): 书签的唯一名称。
    """
    global _bookmark_id_counter
    run = paragraph.runs[0]
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), str(_bookmark_id_counter))
    bookmark_start.set(qn('w:name'), bookmark_name)
    run._r.addprevious(bookmark_start)

    run = paragraph.runs[-1]
    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), str(_bookmark_id_counter))
    run._r.addnext(bookmark_end)

    _bookmark_id_counter += 1


def add_cross_reference_field(paragraph, bookmark_name: str, display_text: str):
    """
    在段落中插入一个带有预显示文本的交叉引用域 (REF field)。

    Args:
        paragraph: python-docx的Paragraph对象。
        bookmark_name (str): 要引用的书签名称。
        display_text (str): 交叉引用在文档中显示的缓存文本。
    """
    run_begin = paragraph.add_run()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run_begin._r.append(fldChar_begin)

    run_instr = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = f' REF {bookmark_name} \\h '
    run_instr._r.append(instrText)

    run_sep = paragraph.add_run()
    fldChar_separate = OxmlElement('w:fldChar')
    fldChar_separate.set(qn('w:fldCharType'), 'separate')
    run_sep._r.append(fldChar_separate)

    run_display = paragraph.add_run()
    run_display.text = display_text

    run_end = paragraph.add_run()
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run_end._r.append(fldChar_end)


def apply_numbering_formats(docx_bytes, page_setup_data):
    """
    通过修改 settings.xml 来应用脚注和尾注的数字格式。

    Args:
        docx_bytes (bytes): 原始DOCX文件字节流。
        page_setup_data (dict): 包含脚注/尾注编号格式的页面设置字典。

    Returns:
        bytes: 修改后的DOCX文件字节流。
    """
    endnote_fmt = page_setup_data.get("endnote_number_format")
    footnote_fmt = page_setup_data.get("footnote_number_format")

    if not endnote_fmt and not footnote_fmt:
        return docx_bytes

    print("\n--- 开始应用自定义引用编号格式 ---")

    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    ET.register_namespace('w', ns['w'])

    temp_dir = tempfile.mkdtemp()
    try:
        with zipfile.ZipFile(io.BytesIO(docx_bytes), 'r') as zip_ref:
            zip_ref.extractall(temp_dir)

        settings_path = os.path.join(temp_dir, 'word', 'settings.xml')

        if not os.path.exists(settings_path):
            print("警告: settings.xml 未找到，跳过编号格式化。")
            return docx_bytes

        settings_tree = ET.parse(settings_path)
        settings_root = settings_tree.getroot()

        if endnote_fmt:
            endnote_pr = settings_root.find('w:endnotePr', ns)
            if endnote_pr is None:
                endnote_pr = ET.SubElement(settings_root, f'{{{ns["w"]}}}endnotePr')
            num_fmt = endnote_pr.find('w:numFmt', ns)
            if num_fmt is None:
                num_fmt = ET.SubElement(endnote_pr, f'{{{ns["w"]}}}numFmt')
            num_fmt.set(f'{{{ns["w"]}}}val', endnote_fmt)
            print(f"  > 已将尾注编号格式设置为: '{endnote_fmt}'")

        if footnote_fmt:
            footnote_pr = settings_root.find('w:footnotePr', ns)
            if footnote_pr is None:
                footnote_pr = ET.SubElement(settings_root, f'{{{ns["w"]}}}footnotePr')
            num_fmt = footnote_pr.find('w:numFmt', ns)
            if num_fmt is None:
                num_fmt = ET.SubElement(footnote_pr, f'{{{ns["w"]}}}numFmt')
            num_fmt.set(f'{{{ns["w"]}}}val', footnote_fmt)
            print(f"  > 已将脚注编号格式设置为: '{footnote_fmt}'")

        settings_tree.write(settings_path, encoding='UTF-8', xml_declaration=True)

        output_stream = io.BytesIO()
        with zipfile.ZipFile(output_stream, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for root, _, files in os.walk(temp_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    archive_name = os.path.relpath(file_path, temp_dir)
                    zipf.write(file_path, archive_name)

        print("--- 自定义引用编号格式应用完毕 ---")
        return output_stream.getvalue()

    finally:
        shutil.rmtree(temp_dir)


def apply_page_setup(doc, page_setup_data: dict):
    """
    根据 page_setup_data 字典中的数据，应用页面布局设置。

    Args:
        doc: python-docx的Document对象。
        page_setup_data (dict): 包含页面设置（方向、边距）的字典。
    """
    if not page_setup_data:
        return

    section = doc.sections[0]

    orientation_str = page_setup_data.get('orientation')
    if orientation_str == 'landscape':
        section.orientation = WD_ORIENT.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height

    margins = page_setup_data.get('margins')
    if isinstance(margins, dict):
        top_cm = margins.get('top')
        if top_cm is not None:
            section.top_margin = Cm(top_cm)
        bottom_cm = margins.get('bottom')
        if bottom_cm is not None:
            section.bottom_margin = Cm(bottom_cm)
        left_cm = margins.get('left')
        if left_cm is not None:
            section.left_margin = Cm(left_cm)
        right_cm = margins.get('right')
        if right_cm is not None:
            section.right_margin = Cm(right_cm)


def apply_section_properties(section, section_data: dict):
    """
    根据 section_data 字典中的数据，应用节的属性（如分栏）。

    Args:
        section: python-docx的Section对象。
        section_data (dict): 包含节属性的字典。
    """
    properties = section_data.get('properties', {})
    if not properties:
        return

    sectPr = section._sectPr
    existing_cols = sectPr.find(qn('w:cols'))
    if existing_cols is not None:
        sectPr.remove(existing_cols)

    columns = properties.get('columns')
    if columns and columns > 1:
        cols = OxmlElement('w:cols')
        cols.set(qn('w:num'), str(columns))
        sectPr.append(cols)


def add_page_break_from_data(doc, element: dict):
    """
    在文档中添加一个分页符。

    Args:
        doc: python-docx的Document对象。
        element (dict): 包含分页符信息的字典（通常为空）。
    """
    doc.add_page_break()


def add_column_break_from_data(doc, element: dict):
    """
    在文档中添加一个分栏符。

    Args:
        doc: python-docx的Document对象。
        element (dict): 包含分栏符信息的字典（通常为空）。
    """
    if not doc.paragraphs:
        p = doc.add_paragraph()
        p.add_run().add_break(WD_BREAK.COLUMN)
        return

    last_paragraph = doc.paragraphs[-1]
    last_paragraph.add_run().add_break(WD_BREAK.COLUMN)


def add_toc_from_data(doc, element: dict):
    """
    在文档中添加一个目录（TOC）。

    Args:
        doc: python-docx的Document对象。
        element (dict): 包含目录信息的字典。
    """
    properties = element.get("properties", {})
    title = properties.get('title')
    if title:
        doc.add_paragraph(title, style='Heading 1')

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()

    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = r'TOC \o "1-3" \h \z \u'
    fldChar_separate = OxmlElement('w:fldChar')
    fldChar_separate.set(qn('w:fldCharType'), 'separate')
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_separate)
    run._r.append(fldChar_end)


async def get_formula_xml_and_placeholder(element: dict) -> tuple[str | None, str | None]:
    """
    将LaTeX公式转换为OMML XML，采用本地转换器优先、LLM兜底的策略。
    增加了对输入字符串的净化，并修正了对异步LLM函数的调用方式。

    Args:
        element (dict): 包含LaTeX公式文本的formula元素字典。

    Returns:
        tuple[str | None, str | None]:
            - str: 用于在文档中临时占位的唯一字符串。
            - str: 转换后的OMML XML字符串或错误提示XML。
    """
    latex_text = element.get("properties", {}).get("text")
    if not latex_text: return None, None

    # 【核心修正 1】净化输入字符串，移除ASCII控制字符
    control_char_re = re.compile(r'[\x00-\x1f\x7f-\x9f]')
    latex_text = control_char_re.sub('', latex_text)

    placeholder = f"__FORMULA_{uuid.uuid4().hex}__"
    omml_str = None

    print(f"⚙️ 尝试使用本地解析器处理: {latex_text}")
    try:
        omml_elem = latex_to_omml(latex_text)
        if omml_elem is not None:
            temp_str = etree.tostring(omml_elem, encoding='unicode')
            etree.fromstring(temp_str)
            if omml_elem.tag == qn('m:oMathPara'):
                print("✅ 本地解析器成功并验证通过。")
                omml_str = temp_str
    except Exception as e:
        print(f"⚠️ 本地解析器失败: {e}。回退到 LLM...")
        omml_str = None

    if not omml_str:
        # 【核心修正 2】正确地 await 异步函数
        llm_xml = await translate_latex_to_omml_llm(latex_text)
        if llm_xml:
            try:
                root = etree.fromstring(llm_xml)
                if root.tag == qn('m:oMathPara') and root.find(qn('m:oMath')) is not None:
                    print("✅ LLM成功并验证通过。")
                    omml_str = llm_xml
                else:
                    print("❌ LLM返回的XML结构不符合规范 (缺少<m:oMathPara>或<m:oMath>)。")
            except etree.XMLSyntaxError as e:
                print(f"❌ LLM返回的不是有效的XML: {e}")

    if omml_str:
        return placeholder, omml_str
    else:
        print(f"❌ 所有转换方法均失败，生成错误提示: '{latex_text}'")
        p = OxmlElement('w:p')
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        color = OxmlElement('w:color');
        color.set(qn('w:val'), 'FF0000')
        rPr.append(color)
        r.append(rPr)
        t = OxmlElement('w:t');
        t.text = f"[公式渲染失败: '{latex_text}']"
        r.append(t)
        p.append(r)
        return placeholder, etree.tostring(p, encoding='unicode')


def post_process_numbering(docx_bytes, numbering_definitions):
    """
    Injects custom multi-level numbering definitions into the .docx file
    by manipulating numbering.xml and styles.xml.
    """
    if not numbering_definitions:
        return docx_bytes

    print("\n--- 开始自定义标题编号XML后处理阶段 ---")

    temp_dir = tempfile.mkdtemp()
    try:
        with zipfile.ZipFile(io.BytesIO(docx_bytes), 'r') as zip_ref:
            zip_ref.extractall(temp_dir)

        numbering_path = os.path.join(temp_dir, 'word', 'numbering.xml')
        styles_path = os.path.join(temp_dir, 'word', 'styles.xml')

        parser = etree.XMLParser(remove_blank_text=True)
        numbering_tree = etree.parse(numbering_path, parser)
        numbering_root = numbering_tree.getroot()
        styles_tree = etree.parse(styles_path, parser)
        styles_root = styles_tree.getroot()

        # Step 1: Create the definitions in numbering.xml
        num_id_map = numbering_generator.create_numbering_definitions(
            numbering_root, numbering_definitions
        )
        print(f"  > 已创建编号定义，映射关系: {num_id_map}")

        # Step 2: Link the styles to the definitions in styles.xml
        numbering_generator.link_styles_to_numbering(
            styles_root, numbering_definitions, num_id_map
        )
        print(f"  > 已将样式链接到编号定义。")

        # Write the modified XML back
        numbering_tree.write(numbering_path, encoding='UTF-8', xml_declaration=True, pretty_print=True)
        styles_tree.write(styles_path, encoding='UTF-8', xml_declaration=True, pretty_print=True)

        # Re-zip the package
        output_stream = io.BytesIO()
        with zipfile.ZipFile(output_stream, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for root, _, files in os.walk(temp_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    archive_name = os.path.relpath(file_path, temp_dir)
                    zipf.write(file_path, archive_name)

        print("--- 自定义标题编号XML后处理阶段完成 ---")
        return output_stream.getvalue()

    finally:
        shutil.rmtree(temp_dir)


async def create_document(data: dict) -> tuple[bytes | None, str | None]:
    """
    【v3 - 完整版】根据经过验证的文档JSON数据创建Word文档。
    此版本是异步的，以处理潜在的异步操作（当前未使用，但为未来保留）。
    它实现了完整的混合内容（run-level）段落渲染，并集成了多个XML后处理步骤，
    包括脚注、尾注、自定义引用格式以及最新的多级列表编号。

    Args:
        data (dict): 包含完整文档结构（页面设置、节、元素）的字典。

    Returns:
        tuple[bytes | None, str | None]:
            - bytes: 最终生成的Word文档的字节流。
            - str: 用于诊断的最终文档body部分的XML字符串。
    """
    doc = Document()

    # --- 初始化 ---
    footnotes_to_inject = {}
    endnotes_to_inject = {}

    # --- 1. 应用全局设置 ---
    page_setup_data = data.get('page_setup', {})
    if page_setup_data:
        apply_page_setup(doc, page_setup_data)

    sections_data = data.get('sections', [])
    if not sections_data:
        sections_data.append({'elements': []})

    # --- 2. 迭代节和元素以构建文档主体 ---
    for i, section_data in enumerate(sections_data):
        section = doc.sections[0] if i == 0 else doc.add_section(WD_SECTION_START.CONTINUOUS)
        apply_section_properties(section, section_data)
        elements = section_data.get('elements', [])

        # 预扫描书签以便交叉引用可以解析显示文本
        bookmark_map = {}
        print("\n--- 开始预扫描本节的书签 ---")
        for element in elements:
            if element.get('type') == 'paragraph':
                properties = element.get('properties', {})
                if properties:
                    bookmark_id = properties.get('bookmark_id')
                    if bookmark_id:
                        text_content = element.get('text', '') or "".join(
                            r.get('text', '') for r in element.get('content', []) if r.get('type') == 'text'
                        )
                        bookmark_map[bookmark_id] = text_content
                        print(f"  > 发现书签: '{bookmark_id}' -> '{text_content[:50]}...'")
        print("--- 预扫描结束 ---\n")

        # 主要元素渲染循环
        for element in elements:
            element_type = element.get('type')

            if element_type == "header":
                add_header_from_data(doc, element, section)
            elif element_type == "footer":
                add_footer_from_data(doc, element, section)
            elif element_type == 'paragraph':
                properties = element.get('properties', {}) or {}
                style = properties.get("style")

                final_style = None
                if style:
                    if any(s.name == style for s in doc.styles):
                        final_style = style
                    else:
                        print(f"警告：找不到名为 '{style}' 的样式，已忽略样式设置。")

                p = doc.add_paragraph(style=final_style)

                # 处理混合内容（run-level）
                content = element.get('content')
                if content and isinstance(content, list):
                    for run_item in content:
                        run_type = run_item.get('type')
                        if run_type == 'text':
                            p.add_run(run_item.get('text', ''))

                        elif run_type == 'formula':
                            latex_text = run_item.get('text', '')
                            if latex_text:
                                omml_element = latex_to_omml(latex_text)
                                if omml_element is not None:
                                    oMath_element = omml_element.find('.//' + qn('m:oMath'))
                                    if oMath_element is not None:
                                        p._p.append(oMath_element)
                                else:
                                    p.add_run(f"[公式渲染失败: {latex_text}]").font.color.rgb = RGBColor(255, 0, 0)

                        elif run_type == 'footnote':
                            placeholder = f"__FOOTNOTE_{uuid.uuid4().hex}__"
                            p.add_run(placeholder)
                            footnotes_to_inject[placeholder] = run_item.get('text', '')

                        elif run_type == 'endnote':
                            placeholder = f"__ENDNOTE_{uuid.uuid4().hex}__"
                            p.add_run(placeholder)
                            endnotes_to_inject[placeholder] = run_item.get('text', '')

                        elif run_type == 'cross_reference':
                            target_bookmark = run_item.get('target_bookmark')
                            if target_bookmark:
                                clean_target = target_bookmark.strip()
                                display_text = bookmark_map.get(clean_target, f"[引用 '{clean_target}' 未找到]")
                                add_cross_reference_field(p, clean_target, display_text)
                                print(f"  > 成功插入对书签 '{clean_target}' 的交叉引用。")

                elif element.get('text'):
                    p.add_run(element.get('text', ''))

                apply_paragraph_properties(p, properties)

                if properties:
                    bookmark_name = properties.get("bookmark_id")
                    if bookmark_name and p.runs:
                        add_bookmark(p, bookmark_name.strip())
                        print(f"> 成功为段落创建书签: '{bookmark_name.strip()}'")

            elif element_type == "table":
                add_table_from_data(doc, element)
            elif element_type == "list":
                add_list_from_data(doc, element)
            elif element_type == "image":
                add_image_from_data(doc, element)
            elif element_type == "page_break":
                add_page_break_from_data(doc, element)
            elif element_type == "column_break":
                add_column_break_from_data(doc, element)
            elif element_type == "toc":
                add_toc_from_data(doc, element)

    # --- 3. 保存到内存流并开始后处理 ---
    draft_stream = io.BytesIO()
    doc.save(draft_stream)
    draft_stream.seek(0)
    processed_bytes = draft_stream.getvalue()
    final_xml_body_for_log = ""

    # --- 4. 执行一系列XML后处理步骤 ---
    page_setup_data_for_refs = data.get('page_setup', {})
    footnote_ref_format = page_setup_data_for_refs.get('footnote_reference_format', "[#]")
    endnote_ref_format = page_setup_data_for_refs.get('endnote_reference_format', "[#]")

    if footnotes_to_inject:
        processed_bytes = post_process_footnotes(processed_bytes, footnotes_to_inject, footnote_ref_format)

    if endnotes_to_inject:
        processed_bytes = post_process_endnotes(processed_bytes, endnotes_to_inject, endnote_ref_format)

    if page_setup_data_for_refs.get("endnote_number_format") or page_setup_data_for_refs.get("footnote_number_format"):
        processed_bytes = apply_numbering_formats(processed_bytes, page_setup_data_for_refs)

    # --- [NEW] 调用多级列表编号后处理器 ---
    numbering_definitions = data.get('numbering_definitions')
    if numbering_definitions:
        processed_bytes = post_process_numbering(processed_bytes, numbering_definitions)

    # --- 5. 生成最终的XML日志以供诊断 ---
    # We extract the final XML body *after* all modifications for the most accurate log.
    try:
        with zipfile.ZipFile(io.BytesIO(processed_bytes), 'r') as zin:
            if 'word/document.xml' in zin.namelist():
                doc_xml = zin.read('word/document.xml')
                root = etree.fromstring(doc_xml)
                body = root.find(qn('w:body'))
                if body is not None:
                    final_xml_body_for_log = etree.tostring(body, pretty_print=True, encoding='unicode')
    except Exception as e:
        print(f"警告: 无法为日志生成最终的XML body: {e}")
        final_xml_body_for_log = "无法为日志生成XML。"

    return processed_bytes, final_xml_body_for_log