# src/latex_converter.py
import re
import traceback
from lxml import etree
from typing import List, Optional, Dict

# --- 1. OMML 命名空间和常量 ---
M_NAMESPACE = "http://schemas.openxmlformats.org/officeDocument/2006/math"
W_NAMESPACE = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M_PREFIX = "{%s}" % M_NAMESPACE
W_PREFIX = "{%s}" % W_NAMESPACE


def _m_tag(tag_name: str) -> str: return M_PREFIX + tag_name


# --- Symbol and Function Maps ---
GREEK_LETTERS = {'\\alpha': 'α', '\\beta': 'β', '\\gamma': 'γ', '\\delta': 'δ', '\\epsilon': 'ε', '\\zeta': 'ζ',
                 '\\eta': 'η', '\\theta': 'θ', '\\iota': 'ι', '\\kappa': 'κ', '\\lambda': 'λ', '\\mu': 'μ', '\\nu': 'ν',
                 '\\xi': 'ξ', '\\omicron': 'ο', '\\pi': 'π', '\\rho': 'ρ', '\\sigma': 'σ', '\\tau': 'τ',
                 '\\upsilon': 'υ', '\\phi': 'φ', '\\chi': 'χ', '\\psi': 'ψ', '\\omega': 'ω', '\\Gamma': 'Γ',
                 '\\Delta': 'Δ', '\\Theta': 'Θ', '\\Lambda': 'Λ', '\\Xi': 'Ξ', '\\Pi': 'Π', '\\Sigma': 'Σ',
                 '\\Upsilon': 'Υ', '\\Phi': 'Φ', '\\Psi': 'Ψ', '\\Omega': 'Ω', '\\varepsilon': 'ɛ', '\\vartheta': 'ϑ',
                 '\\varpi': 'ϖ', '\\varrho': 'ϱ', '\\varsigma': 'ς', '\\varphi': 'ϕ'}
OPERATORS = {'\\pm': '±', '\\times': '×', '\\div': '÷', '\\cdot': '⋅', '\\ast': '∗', '\\cup': '∪', '\\cap': '∩',
             '\\in': '∈', '\\notin': '∉', '\\subset': '⊂', '\\supset': '⊃', '\\subseteq': '⊆', '\\supseteq': '⊇',
             '\\neq': '≠', '\\equiv': '≡', '\\approx': '≈', '\\le': '≤', '\\ge': '≥', '\\geq': '≥', '\\ll': '≪',
             '\\gg': '≫', '\\infty': '∞', '\\nabla': '∇', '\\partial': '∂', '\\forall': '∀', '\\exists': '∃',
             '\\angle': '∠', '\\hbar': 'ħ', '\\prime': '′', '\\leftarrow': '←', '\\rightarrow': '→', '\\to': '→',
             '\\uparrow': '↑', '\\downarrow': '↓', '\\leftrightarrow': '↔', '\\Leftarrow': '⇐', '\\Rightarrow': '⇒',
             '\\implies': '⇒', '\\Uparrow': '⇑', '\\Downarrow': '⇓', '\\Leftrightarrow': '⇔'}
SYMBOLS = {'\\langle': '⟨', '\\rangle': '⟩', '\\{': '{', '\\}': '}', '\\ldots': '…', '\\cdots': '⋯', '\\ddots': '⋱',
           '\\ ': ' ', '\\quad': '    ', '\\,': '\u2009'}
SYMBOL_MAP = {**GREEK_LETTERS, **OPERATORS, **SYMBOLS}
KNOWN_FUNCTIONS = {'\\sin', '\\cos', '\\tan', '\\csc', '\\sec', '\\cot', '\\sinh', '\\cosh', '\\tanh', '\\coth',
                   '\\arcsin', '\\arccos', '\\arctan', '\\log', '\\ln', '\\exp', '\\det', '\\dim', '\\min', '\\max',
                   '\\sup', '\\inf'}
NARY_OPERATORS = {'\\sum': '∑', '\\int': '∫', '\\prod': '∏', '\\oint': '∮', '\\iint': '∬'}
BOUNDARY_TOKENS = ['+', '-', '*', '/', '=', '\\neq', '<', '>', '\\le', '\\ge', ',', '&', '\\\\', '\\right']


def _extract_text(elements: List[etree._Element]) -> str:
    """Recursively extracts all text from a list of OMML elements."""
    text_parts = []
    for elem in elements:
        # Find all text nodes within this element using the correct namespace
        for t in elem.xpath('.//m:t', namespaces={'m': M_NAMESPACE}):
            if t.text:
                text_parts.append(t.text)
    return "".join(text_parts)


# --- 2. OMML 元素构建器 (Element Builders) ---
def _create_run_omml(text: str, is_text: bool = False) -> etree._Element:
    mr = etree.Element(_m_tag('r'))
    if is_text:
        rpr = etree.SubElement(mr, _m_tag('rPr'));
        sty = etree.SubElement(rpr, _m_tag('sty'));
        sty.set(_m_tag('val'), 't')
    mt = etree.SubElement(mr, _m_tag('t'))
    if text.startswith(' ') or text.endswith(' '): mt.set(W_PREFIX + 'space', 'preserve')
    mt.text = text
    return mr


def _create_fraction_omml(num: List[etree._Element], den: List[etree._Element]) -> etree._Element:
    mf = etree.Element(_m_tag('f'));
    mnum = etree.SubElement(mf, _m_tag('num'));
    mden = etree.SubElement(mf, _m_tag('den'))
    for el in num: mnum.append(el)
    for el in den: mden.append(el)
    return mf


def _create_sqrt_omml(base: List[etree._Element]) -> etree._Element:
    mrad = etree.Element(_m_tag('rad'));
    mradPr = etree.SubElement(mrad, _m_tag('radPr'))
    mdegHide = etree.SubElement(mradPr, _m_tag('degHide'));
    mdegHide.set(_m_tag('val'), '1')
    etree.SubElement(mrad, _m_tag('deg'));
    me = etree.SubElement(mrad, _m_tag('e'))
    for el in base: me.append(el)
    return mrad


def _create_accent_omml(base: List[etree._Element], char: str) -> etree._Element:
    macc = etree.Element(_m_tag('acc'));
    maccPr = etree.SubElement(macc, _m_tag('accPr'))
    mchr = etree.SubElement(maccPr, _m_tag('chr'));
    mchr.set(_m_tag('val'), char)
    me = etree.SubElement(macc, _m_tag('e'))
    for el in base: me.append(el)
    return macc


def _create_superscript_omml(base: List[etree._Element], sup: List[etree._Element]) -> etree._Element:
    msSup = etree.Element(_m_tag('sSup'));
    me = etree.SubElement(msSup, _m_tag('e'));
    msup = etree.SubElement(msSup, _m_tag('sup'))
    for el in base: me.append(el)
    for el in sup: msup.append(el)
    return msSup


def _create_subscript_omml(base: List[etree._Element], sub: List[etree._Element]) -> etree._Element:
    msSub = etree.Element(_m_tag('sSub'));
    me = etree.SubElement(msSub, _m_tag('e'));
    msub = etree.SubElement(msSub, _m_tag('sub'))
    for el in base: me.append(el)
    for el in sub: msub.append(el)
    return msSub


def _create_subsup_omml(base: List[etree._Element], sub: List[etree._Element],
                        sup: List[etree._Element]) -> etree._Element:
    msSubSup = etree.Element(_m_tag('sSubSup'));
    me = etree.SubElement(msSubSup, _m_tag('e'))
    msub = etree.SubElement(msSubSup, _m_tag('sub'));
    msup = etree.SubElement(msSubSup, _m_tag('sup'))
    for el in base: me.append(el)
    for el in sub: msub.append(el)
    for el in sup: msup.append(el)
    return msSubSup


def _create_nary_omml(op: str, sub: List[etree._Element], sup: List[etree._Element],
                      base: List[etree._Element]) -> etree._Element:
    mnary = etree.Element(_m_tag('nary'));
    mnaryPr = etree.SubElement(mnary, _m_tag('naryPr'))
    mchr = etree.SubElement(mnaryPr, _m_tag('chr'));
    mchr.set(_m_tag('val'), op)
    if sub:
        msub = etree.SubElement(mnary, _m_tag('sub'))
        for el in sub: msub.append(el)
    if sup:
        msup = etree.SubElement(mnary, _m_tag('sup'))
        for el in sup: msup.append(el)
    me = etree.SubElement(mnary, _m_tag('e'))
    for el in base: me.append(el)
    return mnary


def _create_function_omml(name: str, args: List[etree._Element]) -> etree._Element:
    if args and not args[0].tag == _m_tag('d'): args = [_create_delimiter_omml('(', ')', args)]
    mfunc = etree.Element(_m_tag('func'));
    mfuncName = etree.SubElement(mfunc, _m_tag('fName'))
    mfuncName.append(_create_run_omml(name));
    me = etree.SubElement(mfunc, _m_tag('e'))
    for el in args: me.append(el)
    return mfunc


def _create_delimiter_omml(open_c: str, close_c: str, content: List[etree._Element]) -> etree._Element:
    md = etree.Element(_m_tag('d'));
    mdPr = etree.SubElement(md, _m_tag('dPr'))
    mbeg = etree.SubElement(mdPr, _m_tag('begChr'));
    mbeg.set(_m_tag('val'), open_c)
    mend = etree.SubElement(mdPr, _m_tag('endChr'));
    mend.set(_m_tag('val'), close_c)
    me = etree.SubElement(md, _m_tag('e'))
    for el in content: me.append(el)
    return md


def _create_matrix_omml(rows: List[List[List[etree._Element]]]) -> etree._Element:
    mm = etree.Element(_m_tag('m'))
    for row_data in rows:
        mmr = etree.SubElement(mm, _m_tag('mr'))
        for cell in row_data:
            me = etree.SubElement(mmr, _m_tag('e'))
            if cell:
                for el in cell: me.append(el)
    return mm


def _apply_style_recursively(elements: List[etree._Element], style: str) -> List[etree._Element]:
    for element in elements:
        if element.tag == _m_tag('r'):
            rPr = element.find(f'./{_m_tag("rPr")}')
            if rPr is None: rPr = etree.Element(_m_tag('rPr')); element.insert(0, rPr)
            sty = rPr.find(f'./{_m_tag("sty")}')
            if sty is None: sty = etree.SubElement(rPr, _m_tag('sty'))
            sty.set(_m_tag('val'), style)
        else:
            _apply_style_recursively(list(element), style)
    return elements


# --- 3. LaTeX Parser ---
class ParserState:
    def __init__(self, tokens: List[str]): self.tokens, self.pos = tokens, 0
    def has_tokens(self) -> bool: return self.pos < len(self.tokens)
    def current_token(self) -> Optional[str]: return self.tokens[self.pos] if self.has_tokens() else None
    def advance(self) -> Optional[str]:
        if self.has_tokens():
            token = self.current_token(); self.pos += 1; return token
        return None

def tokenize(latex_string: str) -> List[str]:
    token_regex = re.compile(
        r"(\\begin\{[a-zA-Z]+\*?\}|\\end\{[a-zA-Z]+\*?\}|\\(?:[a-zA-Z]+\*?|\S))"
    )
    latex_parts = token_regex.split(latex_string)
    tokens = []
    for part in latex_parts:
        if not part: continue
        if token_regex.match(part):
            tokens.append(part)
        else:
            tokens.extend(re.findall(r"[{}[\]()|^_&]|[a-zA-Z0-9]+|.", part))
    return [token for token in tokens if not token.isspace()]

def _parse_atomic_expression(state: ParserState) -> List[etree._Element]:
    if state.current_token() == '{':
        state.advance()
        return _parse_tokens(state, stop_tokens=['}'])
    return _parse_single_element(state)

def _parse_argument(state: ParserState) -> List[etree._Element]:
    """
    Parses an argument, which can be a braced group {...}, a parenthesized
    group (...), or a single atomic expression.
    """
    # 1. Handle braced groups (existing logic, works perfectly)
    if state.current_token() == '{':
        state.advance()  # Consume '{'
        return _parse_tokens(state, stop_tokens=['}'])

    # 2. NEW: Handle parenthesized groups intelligently
    if state.current_token() == '(':
        state.advance()  # Consume '('
        # 解析括号内的内容
        content = _parse_tokens(state, stop_tokens=[')'])
        # 返回一个正确的定界符对象，这是修复的关键
        return [_create_delimiter_omml('(', ')', content)]

    # 3. Fallback for single-token arguments (e.g., in `\sqrt x`)
    return _parse_atomic_expression(state)

def _parse_matrix_environment(state: ParserState, env_name: str) -> List[List[List[etree._Element]]]:
    rows, current_row = [], []
    stop_token = f'\\end{{{env_name}}}'
    while state.has_tokens() and state.current_token() != stop_token:
        cell = _parse_tokens(state, stop_tokens=['&', '\\\\', stop_token])
        current_row.append(cell)
        token = state.current_token()
        if token == '&': state.advance()
        elif token == '\\\\':
            rows.append(current_row); current_row = []; state.advance()
        elif token == stop_token: break
    if current_row: rows.append(current_row)
    if state.current_token() == stop_token: state.advance()
    return rows

def _handle_scripts(state: ParserState, base: List[etree._Element]) -> etree._Element:
    sub_arg, sup_arg = None, None
    op1 = state.advance()
    if op1 == '_': sub_arg = _parse_argument(state)
    else: sup_arg = _parse_argument(state)
    if state.current_token() == '_' and sup_arg:
        state.advance(); sub_arg = _parse_argument(state)
    elif state.current_token() == '^' and sub_arg:
        state.advance(); sup_arg = _parse_argument(state)
    if sub_arg and sup_arg: return _create_subsup_omml(base, sub_arg, sup_arg)
    elif sub_arg: return _create_subscript_omml(base, sub_arg)
    return _create_superscript_omml(base, sup_arg)

def _parse_single_element(state: ParserState) -> List[etree._Element]:
    token = state.advance()
    if token.startswith('\\'):
        if token in SYMBOL_MAP: return [_create_run_omml(SYMBOL_MAP[token])]
        if token in KNOWN_FUNCTIONS:
            arg_elements = _parse_argument(state)
            if state.current_token() in ['^', '_']:
                func_el = _create_function_omml(token[1:] + " ", arg_elements)
                return [_handle_scripts(state, [func_el])]
            return [_create_function_omml(token[1:] + " ", arg_elements)]
        if token.startswith('\\operatorname'):
            is_nary_style = token.endswith('*')
            name_elements = _parse_argument(state)
            func_name_text = _extract_text(name_elements)
            if is_nary_style:
                return [_create_run_omml(func_name_text)]
            else:
                arg_elements = _parse_argument(state)
                return [_create_function_omml(func_name_text + " ", arg_elements)]
        if token == '\\lim':
            base = [_create_run_omml('lim')]
            if state.current_token() == '_': return [_handle_scripts(state, base)]
            return base
        if token == '\\frac': return [_create_fraction_omml(_parse_argument(state), _parse_argument(state))]
        if token == '\\sqrt': return [_create_sqrt_omml(_parse_argument(state))]
        if token == '\\hat': return [_create_accent_omml(_parse_argument(state), '^')]
        if token == '\\vec': return [_create_accent_omml(_parse_argument(state), '→')]
        if token == '\\dot': return [_create_accent_omml(_parse_argument(state), '˙')]
        if token == '\\mathbf': return _apply_style_recursively(_parse_argument(state), 'b')
        if token == '\\mathcal': return _parse_argument(state)
        if token == '\\text':
            if state.current_token() == '{':
                state.advance(); raw_text = state.advance(); state.advance()
                return [_create_run_omml(raw_text, is_text=True)]
        if token in NARY_OPERATORS:
            op = NARY_OPERATORS[token]; sub, sup = [], []
            if state.current_token() == '_': state.advance(); sub = _parse_argument(state)
            if state.current_token() == '^': state.advance(); sup = _parse_argument(state)
            base = _parse_tokens(state, stop_tokens=BOUNDARY_TOKENS)
            return [_create_nary_omml(op, sub, sup, base)]
        if token == '\\left':
            open_c = state.advance(); content = _parse_tokens(state, stop_tokens=['\\right'])
            if state.current_token() == '\\right': state.advance()
            close_c = state.advance(); return [_create_delimiter_omml(open_c, close_c, content)]
        if token.startswith('\\begin'):
            env_match = re.match(r'\\begin\{([a-zA-Z]+\*?)\}', token)
            if env_match:
                name = env_match.group(1); mat_el = _create_matrix_omml(_parse_matrix_environment(state, name))
                if name in ['pmatrix', 'pmatrix*']: return [_create_delimiter_omml('(', ')', [mat_el])]
                if name in ['bmatrix', 'bmatrix*']: return [_create_delimiter_omml('[', ']', [mat_el])]
                if name in ['vmatrix', 'vmatrix*']: return [_create_delimiter_omml('|', '|', [mat_el])]
                return [mat_el]
        if token == '\\\\': return [_create_run_omml(' ')]
        print(f"警告: 未知命令 '{token}'"); return [_create_run_omml(token)]
    if token == '{': return _parse_tokens(state, stop_tokens=['}'])
    return [_create_run_omml(token)]

def _parse_tokens(state: ParserState, stop_tokens: List[str] = None) -> List[etree._Element]:
    elements: List[etree._Element] = []
    while state.has_tokens():
        token = state.current_token()
        if stop_tokens and token in stop_tokens:
            if token in ['}', ')', ']']: state.advance()
            break
        base_elements = _parse_single_element(state)
        if state.current_token() in ['^', '_']:
            elements.append(_handle_scripts(state, base_elements))
        else:
            elements.extend(base_elements)
    return elements

def latex_to_omml(latex_string: str, alignment: str = 'center') -> Optional[etree._Element]:
    try:
        tokens = tokenize(latex_string)
        if not tokens: return None
        state = ParserState(tokens); omml_para = etree.Element(_m_tag('oMathPara'))
        omml_para_pr = etree.SubElement(omml_para, _m_tag('oMathParaPr'))
        jc = etree.SubElement(omml_para_pr, _m_tag('jc')); jc.set(_m_tag('val'), alignment)
        omml_math = etree.SubElement(omml_para, _m_tag('oMath'))
        result = _parse_tokens(state)
        for el in result: omml_math.append(el)
        if state.has_tokens(): print(f"警告: 解析在标记 '{state.current_token()}' 处提前结束。")
    except Exception as e:
        print(f"\n{'=' * 25} Critical Error {'=' * 25}\n{type(e).__name__}: {e}");
        traceback.print_exc(); print(f"{'=' * 62}\n"); return None
    return omml_para


# --- 5. 示例和测试 ---
if __name__ == '__main__':
    from docx import Document
    from docx.shared import Pt

    TEST_CASES = [
        r"\oint_{\partial \Omega} \mathbf{F} \cdot d\mathbf{r} = \iint_{\Omega} (\nabla \times \mathbf{F}) \cdot d\mathbf{S}",
        r"\psi(x,t) = \frac{1}{\sqrt{2\pi\hbar}} \int_{-\infty}^{\infty} \phi(p) e^{i(px - Et)/\hbar} \, dp",
        r"\nabla \cdot \mathbf{E} = \frac{\rho}{\varepsilon_0}, \quad \nabla \cdot \mathbf{B} = 0, \quad \nabla \times \mathbf{E} = -\frac{\partial \mathbf{B}}{\partial t}, \quad \nabla \times \mathbf{B} = \mu_0 \mathbf{J} + \mu_0 \varepsilon_0 \frac{\partial \mathbf{E}}{\partial t}",
        r"S = k \ln W = -k \sum_{i} p_i \ln p_i",
        r"\hat{H}\psi = i\hbar \frac{\partial \psi}{\partial t}, \quad \hat{H} = -\frac{\hbar^2}{2m} \nabla^2 + V(\mathbf{r},t)",
        r"\det\begin{pmatrix} E - \epsilon & V \\ V & E - \epsilon \end{pmatrix} = 0 \implies (E - \epsilon)^2 = V^2",
        r"\mathbf{r}(t) = \mathbf{r}_0 \cos(\omega t) + \mathbf{v}_0 \frac{\sin(\omega t)}{\omega} + \frac{\mathbf{F}}{m\omega^2} (1 - \cos(\omega t))",
        r"\Gamma(z) = \int_0^\infty t^{z-1} e^{-t} \, dt, \quad \Gamma(n+1) = n!",
        r"\mathcal{L}\{f(t)\} = \int_0^\infty f(t) e^{-st} \, dt",
        r"\mathbf{A} = \nabla \times \mathbf{F}, \quad \phi = -\nabla \cdot \mathbf{F}",
        r"\frac{d}{dx} \left( x \frac{d y}{dx} \right) + \left( x - \frac{m^2}{x} \right) y = 0 \quad \text{(Bessel equation)}",
        r"\mathbf{T} = \mathbf{I}\alpha + \omega \times (\mathbf{I}\omega)",
        r"\delta S = \int_{t_1}^{t_2} \left[ \frac{\partial L}{\partial q} \delta q + \frac{\partial L}{\partial \dot{q}} \delta \dot{q} \right] dt = 0",
        r"\mathbf{F} = -kx \hat{x} + mg \hat{y}, \quad \omega = \sqrt{\frac{k}{m}}",
        r"\nabla^2 \phi = -\frac{\rho}{\varepsilon_0}, \quad \mathbf{E} = -\nabla \phi",
        r"\int_{-\infty}^{\infty} e^{-a x^2 + b x} \, dx = \sqrt{\frac{\pi}{a}} e^{b^2 / 4a}",
        r"\mathbf{J} = \sigma \mathbf{E} + \rho \mathbf{v} + \nabla \times \mathbf{M}",
        r"\frac{\partial u}{\partial t} = \alpha \frac{\partial^2 u}{\partial x^2}, \quad u(x,0) = f(x)",
        r"\mathbf{P} = \epsilon_0 \chi_e \mathbf{E}, \quad \mathbf{D} = \epsilon_0 \mathbf{E} + \mathbf{P}",
        r"\left[ \hat{x}, \hat{p} \right] = i\hbar, \quad \Delta x \Delta p \geq \frac{\hbar}{2}"
        r"\operatorname*{argmin}_{x \in S} f(x)", r"\operatorname{Tr}(\mathbf{A}) = \sum_i A_{ii}"
    ]


    def main():
        """Generates a DOCX document from a list of LaTeX strings."""
        document = Document()
        document.add_heading('LaTeX to OMML Conversion Test Suite', 0)

        for i, latex_string in enumerate(TEST_CASES):
            print(f"Processing Equation #{i + 1}...")

            document.add_heading(f'Equation #{i + 1}', level=2)
            p = document.add_paragraph()
            p.add_run('Original LaTeX:').bold = True

            code_p = document.add_paragraph(latex_string)
            font = code_p.runs[0].font
            font.name = 'Courier New'
            font.size = Pt(10)

            p = document.add_paragraph()
            p.add_run('Rendered OMML Output:').bold = True

            omml_element = latex_to_omml(latex_string)

            if omml_element is not None:
                # ★★★ THIS IS THE FIX ★★★
                # Create a new paragraph and append the math XML to it.
                # This ensures the document structure is valid.
                p = document.add_paragraph()
                p._p.append(omml_element)
                # ★★★ END OF FIX ★★★
            else:
                document.add_paragraph("--- CONVERSION FAILED ---")

            document.add_paragraph()

        file_path = "latex_test_document.docx"
        document.save(file_path)
        print(f"\nSuccessfully generated document: {file_path}")

    main()

